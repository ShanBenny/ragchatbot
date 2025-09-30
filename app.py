import streamlit as st
import pytesseract
import docx
from PIL import Image
from io import BytesIO

from langchain.schema import Document
from langchain_ollama import OllamaEmbeddings, OllamaLLM
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory


# ================== CONFIG ==================
DOC_PATH = r"C:\Users\Hp\Desktop\ragchatbot\data\C-CARES_User_Manual_VN1.0.docx"
CHROMA_PATH = "vectorstore"


# ================== CUSTOM DOCX LOADER ==================
class DocxWithImagesLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self):
        doc = docx.Document(self.file_path)
        docs = []

        # --- Extract text ---
        text_content = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_content.append(p.text.strip())

        if text_content:
            docs.append(Document(
                page_content="\n".join(text_content),
                metadata={"source": self.file_path, "type": "text"}
            ))

        # --- Extract images ---
        rels = doc.part.rels
        for rel in rels:
            if "image" in rels[rel].target_ref:
                img_bytes = rels[rel]._target.blob
                image = Image.open(BytesIO(img_bytes))

                # OCR text
                img_text = pytesseract.image_to_string(image)

                docs.append(Document(
                    page_content=img_text if img_text.strip() else "[Image with no readable text]",
                    metadata={
                        "source": f"{self.file_path}-image",
                        "type": "image",
                        "image_bytes": img_bytes
                    }
                ))

        return docs


# ================== DATA PIPELINE ==================
def load_and_index():
    loader = DocxWithImagesLoader(DOC_PATH)
    docs = loader.load()

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = splitter.split_documents(docs)

    embeddings = OllamaEmbeddings(model="llama3")
    vectordb = Chroma.from_documents(chunks, embeddings, persist_directory=CHROMA_PATH)
    vectordb.persist()


def get_qa_chain():
    embeddings = OllamaEmbeddings(model="llama3")
    vectordb = Chroma(persist_directory=CHROMA_PATH, embedding_function=embeddings)
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})

    llm = OllamaLLM(model="llama3")

    memory = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True,
        output_key="answer"
    )

    qa = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        memory=memory,
        verbose=True,
        return_source_documents=True,
        output_key="answer"
    )
    return qa


# ================== STREAMLIT UI ==================
st.set_page_config(page_title="C-CARES RAG Chatbot", page_icon="💬", layout="wide")

st.markdown(
    """ 
    <h2 style="margin-bottom:0;">📘 C-CARES RAG Chatbot</h2> 
    <p style="color:gray; margin-top:0;">
    Ask questions about the C-CARES manual (powered by Ollama + LangChain)
    </p> 
    <hr>
    """,
    unsafe_allow_html=True,
)

# Sidebar
with st.sidebar:
    st.header("⚙️ Settings")
    show_sources = st.checkbox("📖 Show Sources", value=True)
    if st.button("🗑️ Clear Chat History"):
        if "qa" in st.session_state:
            st.session_state.qa.memory.clear()
        st.session_state.chat_history = []
        st.success("Chat history cleared!")

    if st.button("🔄 Re-index Document"):
        load_and_index()
        st.success("Document re-indexed!")


# Initialize QA chain
if "qa" not in st.session_state:
    st.session_state.qa = get_qa_chain()
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

qa = st.session_state.qa

# Display previous chat messages
for role, content in st.session_state.chat_history:
    if role == "user":
        st.chat_message("user", avatar="👤").markdown(content)
    else:
        st.chat_message("assistant", avatar="🤖").markdown(content)

# Input box
query = st.chat_input("Type your question about the manual...")

if query:
    st.chat_message("user", avatar="👤").markdown(query)

    result = qa({"question": query})
    answer = result["answer"]

    st.chat_message("assistant", avatar="🤖").markdown(answer)

    # Show sources if enabled
    if show_sources and "source_documents" in result:
        with st.expander("📖 Sources used"):
            for i, doc in enumerate(result["source_documents"], 1):
                st.markdown(f"**Source {i}:** {doc.metadata.get('source', 'Unknown')}")

                if doc.metadata.get("type") == "image" and "image_bytes" in doc.metadata:
                    st.image(doc.metadata["image_bytes"], caption="Image from document", use_column_width=True)

                snippet = doc.page_content.strip()
                if snippet:
                    st.write(snippet[:500] + ("..." if len(snippet) > 500 else ""))

    # Save conversation
    st.session_state.chat_history.append(("user", query))
    st.session_state.chat_history.append(("assistant", answer))
